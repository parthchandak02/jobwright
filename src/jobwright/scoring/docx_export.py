"""Markdown/text-to-DOCX conversion for tailored resumes and cover letters.

Produces editable Word documents for WhatsApp review. Reuses the same
structured-text parser as the PDF pipeline.
"""

from __future__ import annotations

import logging
from pathlib import Path

from jobwright.database import get_connection
from jobwright.scoring.materials_format import normalize_for_structured_parse, resolve_material_path
from jobwright.scoring.pdf import parse_entries, parse_resume, parse_skills

log = logging.getLogger(__name__)


def _add_bullets(doc, text: str) -> None:
    from docx.shared import Pt

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith(("-", "•", "\u2022")):
            stripped = stripped.lstrip("-•\u2022 ").strip()
        p = doc.add_paragraph(stripped, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10)


def _write_structured_docx(text: str, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    resume = parse_resume(normalize_for_structured_parse(text))
    doc = Document()

    if resume.get("name"):
        p = doc.add_paragraph()
        run = p.add_run(resume["name"])
        run.bold = True
        run.font.size = Pt(16)
    if resume.get("title"):
        p = doc.add_paragraph(resume["title"])
        for run in p.runs:
            run.font.size = Pt(11)
    loc_contact = " | ".join(
        x for x in (resume.get("location") or "", resume.get("contact") or "") if x
    )
    if loc_contact:
        p = doc.add_paragraph(loc_contact)
        for run in p.runs:
            run.font.size = Pt(9)

    sections = resume.get("sections") or {}
    if not sections:
        # Plain cover letter / unstructured text
        for para in text.strip().split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    for heading, body in sections.items():
        doc.add_heading(heading.title(), level=2)
        if "SKILL" in heading.upper():
            for cat, val in parse_skills(body):
                p = doc.add_paragraph()
                run = p.add_run(f"{cat}: ")
                run.bold = True
                p.add_run(val)
        elif heading.upper() in ("EXPERIENCE", "PROJECTS", "WORK EXPERIENCE"):
            for entry in parse_entries(body):
                title = entry.get("title") or ""
                subtitle = entry.get("subtitle") or ""
                if title:
                    p = doc.add_paragraph()
                    run = p.add_run(title)
                    run.bold = True
                if subtitle:
                    doc.add_paragraph(subtitle)
                for bullet in entry.get("bullets") or []:
                    bp = doc.add_paragraph(bullet, style="List Bullet")
                    for run in bp.runs:
                        run.font.size = Pt(10)
        else:
            # Summary / education / other: mix of paragraphs and bullets
            if any(line.strip().startswith(("-", "•", "\u2022")) for line in body.splitlines()):
                _add_bullets(doc, body)
            else:
                for para in body.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("DOCX generated: %s", output_path)
    return output_path


def material_to_docx(source_path: Path, output_path: Path | None = None) -> Path:
    """Convert a markdown or legacy text resume/cover letter to DOCX."""
    resolved = resolve_material_path(source_path)
    if not resolved:
        raise FileNotFoundError(f"Material file not found: {source_path}")
    out = Path(output_path) if output_path else resolved.with_suffix(".docx")
    text = resolved.read_text(encoding="utf-8")
    return _write_structured_docx(text, out)


def txt_to_docx(txt_path: Path, output_path: Path | None = None) -> Path:
    """Backward-compatible alias for material_to_docx."""
    return material_to_docx(txt_path, output_path)


def _sibling_docx(material_path: str | None) -> str | None:
    if not material_path:
        return None
    resolved = resolve_material_path(material_path)
    if not resolved:
        return None
    try:
        return str(material_to_docx(resolved))
    except Exception as e:
        log.error("DOCX conversion failed for %s: %s", resolved, e)
        return None


def convert_job_materials(job: dict) -> dict:
    """Convert resume + cover letter txt paths on a job row to DOCX; update DB."""
    resume_docx = _sibling_docx(job.get("tailored_resume_path"))
    cover_docx = _sibling_docx(job.get("cover_letter_path"))
    url = job.get("url")
    if url and (resume_docx or cover_docx):
        conn = get_connection()
        if resume_docx:
            conn.execute(
                "UPDATE jobs SET tailored_resume_docx_path = ? WHERE url = ?",
                (resume_docx, url),
            )
        if cover_docx:
            conn.execute(
                "UPDATE jobs SET cover_letter_docx_path = ? WHERE url = ?",
                (cover_docx, url),
            )
        conn.commit()
    return {
        "url": url,
        "resume_docx": resume_docx,
        "cover_docx": cover_docx,
    }


def batch_convert_docx(limit: int = 50, min_score: int = 5) -> dict:
    """Convert tailored materials to DOCX for eligible jobs; update path columns."""
    conn = get_connection()
    rows = conn.execute(
        """
        SELECT url, title, tailored_resume_path, cover_letter_path,
               tailored_resume_docx_path, cover_letter_docx_path, fit_score
        FROM jobs
        WHERE tailored_resume_path IS NOT NULL
          AND fit_score IS NOT NULL
          AND fit_score >= ?
        ORDER BY fit_score DESC, url
        LIMIT ?
        """,
        (min_score, limit),
    ).fetchall()

    converted = 0
    errors = 0
    for row in rows:
        job = dict(row)
        resume_md = job.get("tailored_resume_path")
        cover_md = job.get("cover_letter_path")
        need_resume = bool(resume_md) and (
            not job.get("tailored_resume_docx_path")
            or not Path(job["tailored_resume_docx_path"]).exists()
        )
        need_cover = bool(cover_md) and (
            not job.get("cover_letter_docx_path")
            or not Path(job["cover_letter_docx_path"]).exists()
        )
        if not need_resume and not need_cover:
            resolved = resolve_material_path(resume_md) if resume_md else None
            if resolved and resolved.with_suffix(".docx").exists() and not job.get(
                "tailored_resume_docx_path"
            ):
                need_resume = True
            elif not need_resume and not need_cover:
                continue

        try:
            convert_job_materials(job)
            converted += 1
        except Exception as e:
            errors += 1
            log.error("DOCX batch failed for %s: %s", job.get("url"), e)

    return {"status": "ok", "converted": converted, "errors": errors}
